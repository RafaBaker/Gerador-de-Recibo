#Bibliotecas para manipulação do Word
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
#Biblioteca para escrever números por extenso
from num2words import num2words

#Definindo a função
def geradorRecibo(inquilino, endereco, tipo, v, mes, vencimento):
    doc = docx.Document('recibo_branco.docx')
    
    num_ext = num2words(v, lang="pt-br")

    #Criando o texto "RECIBO"
    h1 = doc.add_paragraph("RECIBO DE ALUGUEL")
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h1.runs:
        r.font.size = Pt(24)
        r.bold = True
        r.font.name = 'Arial'
        
    space = doc.add_paragraph()
    space.add_run().add_break()
    space.add_run().add_break()


    #Criando o corpo do texto
    p = doc.add_paragraph()
    p.add_run('Recebi do(a) Sr.(a) ')
    p.add_run(f'{inquilino},').bold = True

    p.add_run().add_break()

    p.add_run('A quantia de ')
    p.add_run(f'R$ {v:.2f}').bold = True
    p.add_run(f' ({num_ext}),').bold = True

    p.add_run().add_break()

    p.add_run('Referente ao aluguel do mês de ')
    p.add_run(f'{mes},').bold = True

    p.add_run().add_break()

    p.add_run('Do imóvel de endereço ')
    p.add_run(f'{endereco}, ').bold = True
    p.add_run('de cunho ')
    p.add_run(f'{tipo},').bold = True

    p.add_run().add_break()

    p.add_run('Com vencimento em ')
    p.add_run(f'{vencimento}',).bold = True
    
    p.add_run().add_break()
    p.add_run().add_break()
    
    p.add_run('Emitente responsável: ')
    p.add_run('Rodrigo Moura Matos, ').bold = True
    p.add_run('portador do CPF ')
    p.add_run('046.134.407-69')

    #Espaçamento
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()

    #Campo de assinatura 
    sig = doc.add_paragraph()
    sig.add_run('__________________________________')
    sig.add_run().add_break()
    sig.add_run('Assinatura').bold = True
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Estilizando
    for r in p.runs:
        r.font.size = Pt(14)
        r.font.name = 'Arial'

    for r in sig.runs:
        r.font.size = Pt(14)
        r.font.name = 'Arial'

    #Salvando o documento
    doc.save(f'recibo_{inquilino}_{mes}.docx')
    
    #mensagem de concluído
    print("Recibo gerado!")


#função para o tipo do imóvel
def setTipo(v):
    dic_tipo = {1:"residencial",
                2:"comercial"}
    tipo = dic_tipo[v]
    return tipo

#função para escrever o mês do ano
def setMes(m):
    meses = {1: "janeiro",
             2: "fevereiro",
             3: "março",
             4: "abril",
             5: "maio",
             6: "junho",
             7: "julho",
             8: "agosto",
             9: "setembro",
             10: "outubro",
             11: "novembro",
             12: "dezembro",}
    
    m = meses[m]
    
    return m

#Módulo principal
def main():
    
    #Declarando variáveis
    ctrl = int(1)
    inq = str()
    address = str()
    tipo = int()
    v = int()
    
    #Interface amigável no terminal
    print("="*40)
    string = "GERADOR DE RECIBO!"
    print(" "*11 + string + " "*11)
    print("="*40)
    
    #inicio do loop while
    while ctrl == 1:
        try:
            inq = input("Inquilino: ")
            if not inq.isalpha():
                raise TypeError("Somente letras são permitidas.")
            address = input("Endereco: ")
            tipo = int(input("Tipo [1- Residencial/ 2- Comercial] "))
            v = float(input("Valor: "))
            mes = int(input("Mês [Número]: "))
            vencimento = input("Vencimento: ")
        except:
            print("Valor Inválido! Tente novamente.")
        else:
            #Chamando as funções
            tipo = setTipo(tipo)
            mes = setMes(mes)
            geradorRecibo(inq, address, tipo, v, mes, vencimento)

            p = input("Deseja gerar outro recibo? [s/n] ").lower()
        
            if p == "n":
                ctrl = 0

if __name__ == "__main__":
    main()
